import os
import sys
import re
import json
import shlex
import urllib.parse
import subprocess
import requests
from docx import Document
from config import APP_DIR, UPLOADS_DIR, OLLAMA_URL
import youtube_agent

# --- EXTERNE TOOLS (FFMPEG & YT-DLP) PFADE ---
def get_ffmpeg_path():
	ffmpeg_path = "ffmpeg"
	if getattr(sys, 'frozen', False):
		exe_dir = os.path.dirname(sys.executable)
		local_ffmpeg = os.path.join(exe_dir, "ffmpeg.exe")
		if os.path.exists(local_ffmpeg):
			ffmpeg_path = local_ffmpeg
		else:
			meipass_ffmpeg = os.path.join(sys._MEIPASS, "ffmpeg.exe")
			if os.path.exists(meipass_ffmpeg):
				ffmpeg_path = meipass_ffmpeg
	else:
		local_ffmpeg = os.path.join(APP_DIR, "ffmpeg.exe")
		if os.path.exists(local_ffmpeg):
			ffmpeg_path = local_ffmpeg
	return ffmpeg_path

def get_ytdlp_path():
	if getattr(sys, 'frozen', False):
		exe_dir = os.path.dirname(sys.executable)
		return os.path.join(exe_dir, "yt-dlp.exe")
	else:
		return os.path.join(APP_DIR, "yt-dlp.exe")

def update_ytdlp(is_german=False):
	ytdlp_path = get_ytdlp_path()
	ytdlp_url = "https://github.com/yt-dlp/yt-dlp/releases/latest/download/yt-dlp.exe"
	
	if not os.path.exists(ytdlp_path):
		if is_german:
			print("yt-dlp.exe nicht gefunden. Lade die neueste Version von GitHub herunter...")
		else:
			print("yt-dlp.exe not found. Downloading the latest version from GitHub...")
		try:
			r = requests.get(ytdlp_url, stream=True)
			r.raise_for_status()
			with open(ytdlp_path, 'wb') as f:
				for chunk in r.iter_content(chunk_size=8192):
					f.write(chunk)
			if is_german:
				print("yt-dlp.exe erfolgreich heruntergeladen.")
			else:
				print("yt-dlp.exe successfully downloaded.")
		except Exception as e:
			if is_german:
				print(f"Fehler beim Herunterladen von yt-dlp.exe: {e}")
			else:
				print(f"Error downloading yt-dlp.exe: {e}")
			return
			
	try:
		creationflags = 0x08000000 if os.name == 'nt' else 0
		if is_german:
			print("Prüfe auf yt-dlp Updates...")
		else:
			print("Checking for yt-dlp updates...")
		subprocess.run([ytdlp_path, "--update"], check=False, capture_output=True, creationflags=creationflags)
		if is_german:
			print("yt-dlp Update-Prüfung abgeschlossen.")
		else:
			print("yt-dlp update check completed.")
	except Exception as e:
		if is_german:
			print(f"Fehler beim automatischen Update von yt-dlp: {e}")
		else:
			print(f"Error during automatic update of yt-dlp: {e}")

# --- DOKUMENTEN GENERATOR ---
def process_document_commands(text, de=False):
	matches = re.findall(r'\[SAVE_DOC\](.*?)\[/SAVE_DOC\]', text, re.DOTALL | re.IGNORECASE)
	if not matches:
		return ""
	
	result_text = "\n\n**Dokumenten-Generator:**\n" if de else "\n\n**Document Generator:**\n"
	for match in matches:
		parts = match.split('|||', 1)
		if len(parts) != 2:
			continue
		filename = parts[0].strip()
		content = parts[1].strip()
		
		filename = "".join(c for c in filename if c.isalnum() or c in " ._-")
		if not filename:
			filename = "document.txt"
			
		file_path = os.path.join(UPLOADS_DIR, filename)
		
		try:
			ext = filename.lower().split('.')[-1]
			
			if ext in ['docx', 'doc']:
				doc = Document()
				for line in content.splitlines():
					if line.startswith('# '):
						p = doc.add_heading(level=1)
						line = line[2:]
					elif line.startswith('## '):
						p = doc.add_heading(level=2)
						line = line[3:]
					elif line.startswith('### '):
						p = doc.add_heading(level=3)
						line = line[4:]
					else:
						p = doc.add_paragraph()
					
					parts_line = re.split(r'(\*\*.*?\*\*|__.*?__|\*.*?\*|_.*?_)', line)
					for part in parts_line:
						if (part.startswith('**') and part.endswith('**')) or (part.startswith('__') and part.endswith('__')):
							run = p.add_run(part[2:-2])
							run.bold = True
						elif (part.startswith('*') and part.endswith('*')) or (part.startswith('_') and part.endswith('_')):
							run = p.add_run(part[1:-1])
							run.italic = True
						else:
							p.add_run(part)
				doc.save(file_path)
				
			elif ext == 'rtf':
				rtf_text = content.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
				rtf_text = re.sub(r'\*\*(.*?)\*\*', r'\\b \1\\b0 ', rtf_text)
				rtf_text = re.sub(r'__(.*?)__', r'\\b \1\\b0 ', rtf_text)
				rtf_text = re.sub(r'\*(.*?)\*', r'\\i \1\\i0 ', rtf_text)
				rtf_text = re.sub(r'_(.*?)_', r'\\i \1\\i0 ', rtf_text)
				rtf_text = rtf_text.replace('\n', '\\par\n')
				
				rtf_content = "{\\rtf1\\ansi\\ansicpg1252\\deff0\\nouicompat{\\fonttbl{\\f0\\fnil\\fcharset0 Calibri;}}\n{\\*\\generator ChatUI;}\\viewkind4\\uc1\n\\pard\\sa200\\sl276\\slmult1\\f0\\fs22\\lang7 " + rtf_text + "\n}"
				with open(file_path, "w", encoding="utf-8") as f:
					f.write(rtf_content)
			else:
				with open(file_path, "w", encoding="utf-8") as f:
					f.write(content)
					
			msg = "Datei erfolgreich erstellt" if de else "File successfully created"
			encoded_file = urllib.parse.quote(filename)
			result_text += f"- {msg}: [{filename}](/download/{encoded_file})\n"
		except Exception as e:
			msg = "Fehler beim Erstellen" if de else "Error creating file"
			result_text += f"- {msg}: {str(e)}\n"
			
	return result_text

# --- FFMPEG & YOUTUBE PROCESSING ---
def process_ffmpeg_commands(text, de=False):
	ffmpeg_path = get_ffmpeg_path()
	matches = re.findall(r'\[EXECUTE_FFMPEG\](.*?)(?:\[/EXECUTE_FFMPEG\]|$)', text, re.DOTALL)
	if not matches:
		return ""
	
	result_text = "\n\n**Audio Processing:**\n"
	for match in matches:
		cmd_str = match.strip()
		if not cmd_str:
			continue
		try:
			args = shlex.split(cmd_str)
			if not args:
				continue
			if args[0] != "-i":
				if "-i" in args:
					idx = args.index("-i")
					args = args[idx:]
				else:
					continue
			
			safe_args = []
			for a in args:
				if a == "-y":
					continue
				if "\\" in a or ".." in a:
					raise ValueError(f"Unsafe path detected: {a}")
				if "/" in a:
					# Allow mathematical fractions like 440/432 or 44100*432/440
					# by removing all digit/digit matches and seeing if any slash is left
					temp = re.sub(r'\d+/\d+', '', a)
					if "/" in temp:
						raise ValueError(f"Unsafe path detected: {a}")
				if "://" in a:
					raise ValueError(f"URL schemes not allowed: {a}")
				if a == "concat":
					raise ValueError("Concat demuxer is not allowed for security reasons.")
				safe_args.append(a)
			
			full_cmd = [ffmpeg_path, "-y"] + safe_args
			output_file = safe_args[-1]
			
			creationflags = 0x08000000 if os.name == 'nt' else 0
			subprocess.run(full_cmd, cwd=UPLOADS_DIR, check=True, capture_output=True, creationflags=creationflags)
			msg = "Erfolgreich verarbeitet" if de else "Successfully processed"
			result_text += f"- {msg}: [{output_file}](/download/{urllib.parse.quote(output_file)})\n"
		except FileNotFoundError:
			result_text += f"- Fehler: ffmpeg.exe was not found. Please place the file next to ChatUI.exe.\n"
		except subprocess.CalledProcessError as e:
			err_out = e.stderr.decode('utf-8', errors='ignore') if e.stderr else str(e)
			msg = "Konvertierungsfehler" if de else "Conversion error"
			result_text += f"- {msg}: {err_out[-200:]}\n"
		except Exception as e:
			msg = "Allgemeiner Fehler" if de else "General error"
			result_text += f"- {msg}: {str(e)}\n"
	return result_text

def process_youtube_commands(text, de=False):
	matches = re.findall(r'\[PLAY_YOUTUBE\](.*?)\[/PLAY_YOUTUBE\]', text, re.DOTALL | re.IGNORECASE)
	if not matches:
		return ""
	
	ffmpeg_path = get_ffmpeg_path()
	ytdlp_path = get_ytdlp_path()
	result_text = "\n\n**YouTube Player:**\n"
	
	for match in matches:
		query = match.strip()
		if not query:
			continue
		
		filename, title_or_error = youtube_agent.search_and_download_audio(query, UPLOADS_DIR, ffmpeg_path, ytdlp_path)
		
		if filename:
			encoded_file = urllib.parse.quote(filename)
			result_text += f"Audio: **{title_or_error}**\n"
			result_text += f"<audio controls autoplay src=\"/media/{encoded_file}\" style=\"width:100%; outline:none; margin-top:10px; border-radius: 8px;\"></audio>\n"
		else:
			msg = "Fehler beim Laden (YouTube)" if de else "Error loading (YouTube)"
			result_text += f"{msg}: {title_or_error}\n"
			
	return result_text

# --- E-MAIL EXTRAKTION & INTENT ---
def extract_email_info(text, model, user_settings):
	prompt = f'Extract the recipient (to) and subject. Then, WRITE a complete email message based on the input for the "body" field. Do not just copy the input. Return ONLY a valid JSON object: {{"to": "...", "subject": "...", "body": "..."}}.\n\nText: {text}'
	try:
		response = ""
		if user_settings.get("ai_provider") == "gemini" and user_settings.get("gemini_api_key"):
			url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={user_settings['gemini_api_key']}"
			payload = {"contents": [{"role": "user", "parts": [{"text": prompt}]}]}
			r = requests.post(url, json=payload, headers={"Content-Type": "application/json"})
			r.raise_for_status()
			j = r.json()
			response = j.get("candidates", [{}])[0].get("content", {}).get("parts", [{}])[0].get("text", "").strip()
		elif user_settings.get("ai_provider") == "openrouter" and user_settings.get("openrouter_api_key"):
			url = "https://openrouter.ai/api/v1/chat/completions"
			headers = {"Authorization": f"Bearer {user_settings['openrouter_api_key']}", "Content-Type": "application/json"}
			payload = {"model": model, "messages": [{"role": "user", "content": prompt}], "stream": False}
			r = requests.post(url, json=payload, headers=headers)
			r.raise_for_status()
			j = r.json()
			response = j.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
		elif user_settings.get("ai_provider") == "openai" and user_settings.get("openai_api_key"):
			url = "https://api.openai.com/v1/chat/completions"
			headers = {"Authorization": f"Bearer {user_settings['openai_api_key']}", "Content-Type": "application/json"}
			payload = {"model": model, "messages": [{"role": "user", "content": prompt}], "stream": False}
			r = requests.post(url, json=payload, headers=headers)
			r.raise_for_status()
			j = r.json()
			response = j.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
		else:
			payload = {"model": model, "prompt": prompt, "stream": False}
			r = requests.post(f"{OLLAMA_URL}/api/generate", json=payload)
			r.raise_for_status()
			response = r.json().get("response", "").strip()

		match = re.search(r'\{.*\}', response, re.DOTALL)
		if match:
			response = match.group(0)

		return json.loads(response)
	except Exception as e:
		print("Fehler bei E-Mail JSON Extraktion:", e)
		return None

def parse_email_intent(text, model, user_settings):
	prompt = f'Analyze the user request regarding reading, searching or deleting emails. Determine if the user wants to READ/SEARCH emails, or DELETE emails. Return ONLY a valid JSON object. For read/search: {{"action": "read", "keyword": "optional search term, otherwise empty"}}. For delete: {{"action": "delete", "uids": ["123", "124"]}}. Text: {text}'
	try:
		response = ""
		if user_settings.get("ai_provider") == "gemini" and user_settings.get("gemini_api_key"):
			url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={user_settings['gemini_api_key']}"
			payload = {"contents": [{"role": "user", "parts": [{"text": prompt}]}]}
			r = requests.post(url, json=payload, headers={"Content-Type": "application/json"})
			r.raise_for_status()
			j = r.json()
			response = j.get("candidates", [{}])[0].get("content", {}).get("parts", [{}])[0].get("text", "").strip()
		elif user_settings.get("ai_provider") == "openrouter" and user_settings.get("openrouter_api_key"):
			url = "https://openrouter.ai/api/v1/chat/completions"
			headers = {"Authorization": f"Bearer {user_settings['openrouter_api_key']}", "Content-Type": "application/json"}
			payload = {"model": model, "messages": [{"role": "user", "content": prompt}], "stream": False}
			r = requests.post(url, json=payload, headers=headers)
			r.raise_for_status()
			j = r.json()
			response = j.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
		elif user_settings.get("ai_provider") == "openai" and user_settings.get("openai_api_key"):
			url = "https://api.openai.com/v1/chat/completions"
			headers = {"Authorization": f"Bearer {user_settings['openai_api_key']}", "Content-Type": "application/json"}
			payload = {"model": model, "messages": [{"role": "user", "content": prompt}], "stream": False}
			r = requests.post(url, json=payload, headers=headers)
			r.raise_for_status()
			j = r.json()
			response = j.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
		else:
			payload = {"model": model, "prompt": prompt, "stream": False}
			r = requests.post(f"{OLLAMA_URL}/api/generate", json=payload)
			r.raise_for_status()
			response = r.json().get("response", "").strip()

		match = re.search(r'\{.*\}', response, re.DOTALL)
		if match:
			response = match.group(0)

		return json.loads(response)
	except Exception as e:
		return {"action": "read", "keyword": ""}